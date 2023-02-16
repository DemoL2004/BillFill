from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words as n
from  datetime import date as dt
import datetime
from functools import reduce
from math import ceil,floor
gauth = GoogleAuth()
drive = GoogleDrive(gauth)
check=int(input("Enter number 1 for Transport bill and number 2 for Breaking bill"))
if check==1:
    d=Document("tbtemplate.docx")
if check==2:
    d=Document("bbtemplate.docx")
if check !=3:
    dic={"1":"31","2":"28","3":"31","4":"30","5":"31","6":"30","7":"31","8":"31","9":"30","10":"31","11":"30","12":"31"}
    todays=dt.today()
    tdf=todays.strftime("%d/%m/%Y")
    tdfl=tdf.split("/")
    if tdfl[1]=="01"or tdfl[1]=="1":
        tdfl[1] = "12"
        tdfl[2] = str(int(tdfl[2]) - 1)

    else:
        tdfl[1] = str(int(tdfl[1]) - 1)
    pmd=dic[tdfl[1]]
    ld=pmd
    ld+="/"
    ld+=tdfl[1]
    ld+="/"
    ld+=tdfl[2]

    sd="01"
    sd+="/"
    sd+=tdfl[1]
    sd+="/"
    sd+=tdfl[2]
    tlis=[tdf,sd,ld]
    dlis=[]
    for f  in tlis:
        tem1=f.split("/")
        tempo=""
        for i in range(len(tem1)):
            if f==sd:
                sdy=tem1[-1]
                sdy=sdy[2:]
                if i!=2:
                    fsd=tem1[i]
                else:
                    fsd=sdy
                tempo += fsd
            else:
                tempo+=tem1[i]
            if i<=1:
                tempo+="."
        dlis.append(tempo)
    for i in range(len(d.paragraphs)):
        #to change invoice no
        if i==2:
            f=open("inv.txt","r")
            a=d.paragraphs[i]
            b=a.runs[-1]
            bo=b.text.split()[-1]
            c=int(f.read())+1
            f.close()
            f=open("inv.txt","w")
            f.write(str(c))
            f.close()
            te=" "
            for i in b.text.split():
                if i !=bo:
                    te+=i
                    te+=" "
            te+=str(c)
            b.text=te
        if i==3:
            a = d.paragraphs[i]
            b = a.runs[-1]
            b.text=dlis[0]
        if i==4:
            a = d.paragraphs[i]
            b = a.runs[-1]
            b.text=f"Period:{dlis[1]} to {dlis[2]}"




    def formatINR(number):
        s, *d = str(number).partition(".")
        r = ",".join([s[x - 2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
        return "".join([r] + d)
    if check==1:
        for i in d.tables:
            btq=input("Enter quantity for Boulder transport: ")
            atq = input("Enter quantity for Aggregates transport: ")
            wtq = input("Enter quantity for Waste transport: ")
            rate=24.99
            i.rows[1].cells[2].text=btq
            p=i.rows[1].cells[2].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount=str(float(btq)*rate)
            amount=round(float(amount),2)
            al=[]
            al.append(amount)

            i.rows[1].cells[4].text = str(formatINR(amount))

            p=i.rows[1].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[2].cells[2].text=atq
            p=i.rows[2].cells[2].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount = str(float(atq) * rate)
            amount = round(float(amount), 2)
            al.append(amount)
            i.rows[2].cells[4].text = str(formatINR(amount))
            p=i.rows[2].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[3].cells[2].text = wtq
            p=i.rows[3].cells[2].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount = str(float(wtq) * rate)
            amount = round(float(amount), 2)
            al.append(amount)
            i.rows[3].cells[4].text =str(formatINR(amount))
            p=i.rows[3].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            def add(a,b):
                return a+b
            total=reduce(add,al)
            total=round(float(total), 2)
            i.rows[4].cells[4].text=str(formatINR(total))
            p=i.rows[4].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            gst=2.5
            amount = str((total*gst)/100)
            amount = round(float(amount), 2)
            al.append(amount)
            al.append(amount)
            i.rows[5].cells[4].text=str(formatINR(amount))
            p=i.rows[5].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[6].cells[4].text = str(formatINR(amount))
            p=i.rows[6].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            total = reduce(add, al)
            total = round(float(total), 2)
            i.rows[7].cells[4].text = str(formatINR(total))
            p=i.rows[7].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount = str(total)
            roundof = ""
            if amount[-2] != ".":
                roundof += amount[-2]
                roundof += amount[-1]
                roundof = int(roundof)
                if roundof>=51:
                    gt=ceil(total)
                else:
                    gt=floor(total)
            else:
                roundof += amount[-1]
                roundof = int(roundof)
                if roundof > 5:
                    gt = ceil(total)
                else:
                    gt = floor(total)
            i.rows[8].cells[4].text = str(formatINR(gt))
            p = i.rows[8].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for i in range(len(d.paragraphs)):
            if i == 14:

                p = n(gt, lang='en_IN')

                p = list(p)
                for iii in range(len(p)):
                    if p[iii] == ",":
                        p[iii] = " and "
                pp = ""
                for iii in p:
                    pp += iii
                p = pp
                pp = p.split()
                for beh in range(len(pp)):
                    if pp[beh] in ["thousand", "and", "hundred", "lakh"]:
                        continue

                    pp[beh] = pp[beh].capitalize()
                p = ""
                for beh in pp:
                    p += beh
                    p += " "
                a = d.paragraphs[i]
                b = a.runs[-1]
                b.text=f"Amount in words: {p} only"
        today = datetime.date.today()
        first = today.replace(day=1)
        last_month = first - datetime.timedelta(days=1)
        lmst=last_month.strftime("%b")
        lmst+=" "
        lmst+=sdy
        d.save(f"Transport bill {lmst}.docx")
        """upload_file = f"Transport bill {lmst}.docx"
        gfile = drive.CreateFile({'parents': [{'id': '1kj9W_HYBSeNv11fA96reI2qo-TFfEj5D'}]})
        # Read file and set it as the content of this instance.
        gfile.SetContentFile(upload_file)
        gfile.Upload()  # Upload the file."""
    if check==2:
        for i in d.tables:
            bq=input("Enter quantity for Breaking: ")
            qlq = input("Enter quantity for Quarry Loading: ")
            slq = input("Enter quantity for Stock Loading: ")
            mlq = input("Enter quantity for Machinery Loading: ")
            bqr=31.99
            qlr=29.99
            slr=30
            mlqr=2100
            i.rows[1].cells[2].text = bq
            p = i.rows[1].cells[2].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[2].cells[2].text = qlq
            p = i.rows[2].cells[2].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[3].cells[2].text = slq
            p = i.rows[3].cells[2].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[4].cells[2].text = mlq
            p = i.rows[4].cells[2].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount = str(float(bq) * bqr)
            amount = round(float(amount), 2)
            al = []
            al.append(amount)
            i.rows[1].cells[4].text = str(formatINR(amount))
            p = i.rows[1].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            def add(a, b):
                return a + b
            amount = str(float(qlq) * qlr)
            amount = round(float(amount), 2)
            al.append(amount)
            i.rows[2].cells[4].text = str(formatINR(amount))
            p = i.rows[2].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount=str(float(slq)*slr)
            amount=round(float(amount),2)
            al.append(amount)
            i.rows[3].cells[4].text = str(formatINR(amount))
            p = i.rows[3].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount = str(float(mlq) * mlqr)
            amount = round(float(amount), 2)
            al.append(amount)
            i.rows[4].cells[4].text = str(formatINR(amount))
            p = i.rows[4].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            total=reduce(add,al)
            total=round(float(total), 2)
            i.rows[5].cells[4].text=str(formatINR(total))
            p=i.rows[5].cells[4].paragraphs[0]
            p.alignment =WD_PARAGRAPH_ALIGNMENT.RIGHT
            gst=9
            amount = str((total*gst)/100)
            amount = round(float(amount), 2)
            al.append(amount)
            al.append(amount)
            i.rows[6].cells[4].text = str(formatINR(amount))
            p = i.rows[6].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            i.rows[7].cells[4].text = str(formatINR(amount))
            p = i.rows[7].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            total = reduce(add, al)
            total = round(float(total), 2)
            i.rows[8].cells[4].text = str(formatINR(total))
            p = i.rows[8].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            amount = str(total)
            roundof = ""
            if amount[-2] != ".":
                roundof += amount[-2]
                roundof += amount[-1]
                roundof = int(roundof)
                if roundof >= 51:
                    gt = ceil(total)
                else:
                    gt = floor(total)
            else:
                roundof += amount[-1]
                roundof = int(roundof)
                if roundof > 5:
                    gt = ceil(total)
                else:
                    gt = floor(total)
            i.rows[9].cells[4].text = str(formatINR(gt))
            p = i.rows[9].cells[4].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for i in range(len(d.paragraphs)):
            if i == 14:

                p = n(gt, lang='en_IN')

                p = list(p)
                for iii in range(len(p)):
                    if p[iii] == ",":
                        p[iii] = " and "
                pp = ""
                for iii in p:
                    pp += iii
                p = pp
                pp = p.split()
                for beh in range(len(pp)):
                    if pp[beh] in ["thousand", "and", "hundred", "lakh"]:
                        continue
                    pp[beh] = pp[beh].capitalize()
                p = ""
                for beh in pp:
                    p += beh
                    p += " "
                a = d.paragraphs[i]
                b = a.runs[-1]
                b.text=f"Amount in words: {p} only"
        today = datetime.date.today()
        first = today.replace(day=1)
        last_month = first - datetime.timedelta(days=1)
        lmst=last_month.strftime("%b")
        lmst += " "
        lmst+=sdy
        d.save(f"Breaking bill {lmst}.docx")
        """upload_file = f"Breaking bill {lmst}.docx"
        gfile = drive.CreateFile({'parents': [{'id': '1kj9W_HYBSeNv11fA96reI2qo-TFfEj5D'}]})
        # Read file and set it as the content of this instance.
        gfile.SetContentFile(upload_file)
        gfile.Upload()  # Upload the file."""
